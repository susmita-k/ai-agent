import os
import base64
from io import BytesIO
from docx import Document
from dotenv import load_dotenv  
from pydantic import BaseModel
from fastapi import FastAPI
from objects import PatientInput
from crewai import Agent, Task, Crew
from langchain_openai import ChatOpenAI
from fastapi.middleware.cors import CORSMiddleware
from crewai_tools import ScrapeWebsiteTool, SerperDevTool

load_dotenv()
app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:8000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Set up keys
# os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")
# os.environ["SERPER_API_KEY"] = os.getenv("SERPER_API_KEY")

os.environ["OPENAI_API_KEY"] = ""
os.environ["SERPER_API_KEY"] = ""

# Tools and LLM
search_tool = SerperDevTool()
scrape_tool = ScrapeWebsiteTool()

llm = ChatOpenAI(
    model="gpt-3.5-turbo-16k",
    temperature=0.1,
    max_tokens=3000
)

# Agents and Tasks
role_designator = Agent(
    role="Role Designator",
    goal="Convert an undifferentiated transcript of a voice conversation between a doctor and a patient into structured chronological dialogue fragments, attributing each line to either the doctor or the patient.",
    backstory=(
        "Trained on thousands of hours of clinical conversations, this agent is a specialist in identifying speaker roles "
        "in healthcare dialogues. With expertise in medical linguistics, pragmatic cues, and contextual reasoning, "
        "it discerns subtle distinctions in speech—whether it's a physician delivering a diagnosis, a nurse asking clarifying questions, "
        "or a patient expressing symptoms or concerns. Its primary mission is to ensure accurate role attribution, "
        "serving as the foundation for high-precision downstream tasks like summarization, speech translation, "
        "medical charting, and intelligent response generation in clinical settings."
    ),
    verbose=True,
    allow_delegation=False,
    tools=[search_tool, scrape_tool],
    llm=llm
)

role_designator_task = Task(
    description=(
        "1. Process a raw voice transcript containing undifferentiated speech from both a doctor and a patient. Transcript is as follows: ({conversation}).\n"
        "2. Accurately differentiate and label the speech fragments as either 'doctor' or 'patient' based on context and language cues.\n"
        "3. Segment the transcript chronologically into a JSON array, where each element contains one fragment of speech from the doctor and one from the patient, paired together in the order they occurred.\n"
        "4. Do not return any intermediate steps or explanations, only the final JSON array.\n"
        "Conversation Input:${inputs['conversation']}\n"
    ),
    expected_output=(
        "A JSON array in the following format, with each element representing a chronological pair of dialogue fragments from the doctor and the patient:\n"
        "[{'doctor': 'doctor speech first fragment','patient': 'patient speech first fragment'},{'doctor': 'doctor speech next fragment', 'patient': 'patient speech next fragment'}]\n"
    ),
    agent=role_designator
)


crew = Crew(
    agents=[role_designator],
    tasks=[role_designator_task],
    verbose=True     
)


# DOCX generation
def generate_docx(result):
    doc = Document()
    doc.add_heading('Healthcare Diagnosis and Treatment Recommendations', 0)
    result = " ".join(str(item) for item in result)
    doc.add_paragraph(result)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


@app.post("/diagnose")
def get_diagnosis(data: PatientInput):
    print("[/diagnose]Received data:", data)
    result = None
    sample_data = data.symptoms
    
    #sample_data = "How are you feeling today? I've had a persistent cough and mild fever. Any shortness of breath or chest pain? No, just fatigue. Sounds viral, but we'll run some tests to be sure. Thank you, doctor."
    try:
        print("Calling crew.kickoff with data - sample_data:", sample_data)
        #result = crew.kickoff(inputs={"symptoms": data.symptoms, "medical_history": data.medical_history})
        #result = crew.kickoff(inputs={"symptoms": sample_data, "medical_history": data.medical_history})
        result = crew.kickoff(inputs={"conversation": sample_data, "medical_history": data.medical_history})
        print("Crew kickoff result:", result)
        docx_file = generate_docx(result)
        docx_bytes = docx_file.read()
    except Exception as e:
        print("Error during crew.kickoff:", e)
        docx_file = generate_docx("ERROR")
        docx_bytes = docx_file.read()
    finally:
        print("Crew kickoff finally.")
        encoded_doc = base64.b64encode(docx_bytes).decode('utf-8')
        return {
            "diagnosis_summary": result,
            "docx_base64": encoded_doc
        }